"""Word 文档文本提取（python-docx）"""
from __future__ import annotations

import os
from typing import Any, Dict, List, Optional

from app.core.logger import log


try:
    import docx  # python-docx
    HAS_DOCX = True
except Exception:  # pragma: no cover
    docx = None  # type: ignore
    HAS_DOCX = False


class DocxExtractor:
    """Word 简历文本提取器"""

    def extract(self, file_path: str) -> Dict[str, Any]:
        if not os.path.exists(file_path):
            return {"text": "", "paragraphs": [], "tables": [], "error": "file_not_found"}
        if not HAS_DOCX:
            return self._mock_extract(file_path)
        try:
            doc = docx.Document(file_path)
            paragraphs: List[str] = [p.text for p in doc.paragraphs if p.text]
            tables: List[List[List[str]]] = []
            for tbl in doc.tables:
                grid: List[List[str]] = []
                for row in tbl.rows:
                    grid.append([cell.text for cell in row.cells])
                tables.append(grid)
            text = "\n".join(paragraphs + [c for row in tables for line in row for c in line])
            return {"text": text, "paragraphs": paragraphs, "tables": tables}
        except Exception as e:
            log.error(f"python-docx 解析失败: {e}")
            return self._mock_extract(file_path)

    def extract_bytes(self, data: bytes) -> Dict[str, Any]:
        if not HAS_DOCX:
            return {"text": "", "paragraphs": [], "tables": [], "error": "docx_unavailable"}
        try:
            import io
            doc = docx.Document(io.BytesIO(data))
            paragraphs = [p.text for p in doc.paragraphs if p.text]
            return {"text": "\n".join(paragraphs), "paragraphs": paragraphs, "tables": []}
        except Exception as e:
            log.error(f"docx 字节流解析失败: {e}")
            return {"text": "", "paragraphs": [], "tables": [], "error": str(e)}

    @staticmethod
    def _mock_extract(file_path: str) -> Dict[str, Any]:
        log.warning("python-docx 未安装, 返回空 docx 结果")
        return {
            "text": f"[MOCK DOCX] {os.path.basename(file_path)}",
            "paragraphs": [],
            "tables": [],
            "mock": True,
        }


_singleton: Optional[DocxExtractor] = None


def get_docx_extractor() -> DocxExtractor:
    global _singleton
    if _singleton is None:
        _singleton = DocxExtractor()
    return _singleton
